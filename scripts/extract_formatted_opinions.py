#!/usr/bin/env python3
"""
Extract formatted court opinions from DOCX files with bidirectional footnotes.
Preserves all formatting (bold, italic, underline, etc.) except hyperlinks.
"""

import sys
import json
import re
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

sys.stdout.reconfigure(encoding='utf-8')


def extract_footnotes(doc: Document) -> Dict[str, str]:
    """Extract footnotes from DOCX with formatting preserved."""
    footnotes_dict = {}

    try:
        footnotes_part = doc.part.part_related_by(RT.FOOTNOTES)
    except KeyError:
        return footnotes_dict

    # Parse footnotes XML
    footnotes_xml = etree.fromstring(footnotes_part.blob)
    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    footnotes = footnotes_xml.findall(f"{w_ns}footnote")

    for fn in footnotes:
        fn_id = fn.get(qn('w:id'))

        # Skip default footnotes 0 and 1
        if fn_id in ('0', '1'):
            continue

        # Extract all paragraphs from footnote
        fn_paras = fn.findall(f"{w_ns}p")
        fn_parts = []

        for para in fn_paras:
            para_html = extract_formatted_paragraph_xml(para, w_ns)
            if para_html.strip():
                fn_parts.append(para_html)

        fn_html = '\n'.join(fn_parts)

        # Remove leading footnote number with any formatting wrapping (e.g., "<p><sup><u><em><strong>9 </strong></em></u></sup>...")
        # This regex removes everything from <p> until the first letter
        fn_html = re.sub(r'(<p>)\s*(?:<[^>]*>)*\s*\d+\s*(?:</[^>]*>)*\s*', r'\1', fn_html)

        footnotes_dict[fn_id] = fn_html

    return footnotes_dict


def extract_formatted_paragraph_xml(element, w_ns: str) -> str:
    """Extract paragraph from XML element with formatting preserved.
    Skip artifact patterns: bold+italic+underline together indicates DOCX creation artifact."""
    html_parts = []

    # Get all text runs
    runs = element.findall(f"{w_ns}r")

    for run in runs:
        # Skip footnote references in main text (handled separately)
        if run.find(f"{w_ns}footnoteReference") is not None:
            continue

        # Get text content
        text_elem = run.find(f"{w_ns}t")
        text = text_elem.text if text_elem is not None and text_elem.text else ""

        if not text:
            continue

        # Strip URLs from text (remove hyperlinks)
        text = re.sub(r'\s*https?://\S+', '', text)

        # Escape HTML special characters
        text = escape_html(text)

        # Extract formatting from run properties
        rpr = run.find(f"{w_ns}rPr")
        if rpr is not None:
            # Check for artifact patterns: if all three formatting flags are set, skip all
            has_b = rpr.find(f"{w_ns}b") is not None
            has_i = rpr.find(f"{w_ns}i") is not None
            has_u = rpr.find(f"{w_ns}u") is not None
            is_artifact = has_b and has_i and has_u

            # Apply formatting selectively
            if not is_artifact:
                # Bold - only if not part of artifact
                if has_b:
                    text = f"<strong>{text}</strong>"

                # Italic - only if not part of artifact (preserve case name italics)
                if has_i:
                    text = f"<em>{text}</em>"

                # Underline - only if not part of artifact
                if has_u:
                    text = f"<u>{text}</u>"

            # Superscript and subscript are always applied (never artifacts)
            va = rpr.find(f"{w_ns}vertAlign")
            if va is not None and va.get(qn('w:val')) == 'superscript':
                text = f"<sup>{text}</sup>"

            if va is not None and va.get(qn('w:val')) == 'subscript':
                text = f"<sub>{text}</sub>"

        html_parts.append(text)

    content = ''.join(html_parts)

    if content.strip():
        return f"<p>{content}</p>"
    return ""


def find_opinion_start(doc: Document) -> int:
    """Find the paragraph where the actual opinion text begins (marked by judge name)."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        # Look for patterns like "Reese, Judge" or "Barnes, Presiding Judge"
        if re.search(r'(Judge|Justice)\s*\.?\s*$', text) and i > 15:
            return i
        # Also check for "[*page]" pattern which marks opinion start
        if re.match(r'^\[\*\d+\]', text) and i > 10:
            return i
    return 20  # Default fallback


def italicize_case_names(text: str) -> str:
    """Add italics to case names in citations. Matches patterns with 'v.' or 'of/Interest' type case names."""
    # Match two patterns:
    # 1. Standard case citations with "v.": "Name v. Name, citation"
    # 2. In the Interest of / Ex parte / etc.: "In the Interest of X, citation"
    pattern1 = r'([A-Z][A-Za-z\.\s&,]*?v\.\s+[A-Z][A-Za-z\.\s&0-9]*?)(?=,\s*\d+\s+|,\s*\[)'
    pattern2 = r'((?:In the Interest of|Ex parte)\s+[A-Za-z\.\s]+?)(?=,\s*\d+\s+|,\s*\[)'

    def replace_case(match):
        case_name = match.group(1).strip()
        # Skip if already italicized
        if '<em>' in case_name or '</em>' in case_name:
            return case_name
        return f'<em>{case_name}</em>'

    # Only apply to text outside of existing em tags
    # Split on em tags to preserve them, process plain text parts only
    parts = re.split(r'(<em>.*?</em>)', text)
    result = []
    for i, part in enumerate(parts):
        if i % 2 == 0:  # Not inside em tags
            # Apply both patterns
            part = re.sub(pattern1, replace_case, part)
            part = re.sub(pattern2, replace_case, part)
        result.append(part)

    return ''.join(result)


def extract_formatted_opinion(doc: Document, footnotes_dict: Dict[str, str]) -> Tuple[str, Dict[str, str]]:
    """Extract opinion text with all formatting and bidirectional footnotes."""
    html_parts = []
    footnote_refs = {}  # Track which footnotes appear in text

    w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    # Find where the actual opinion begins (skip metadata)
    opinion_start = find_opinion_start(doc)

    for para in doc.paragraphs[opinion_start:]:
        # Skip empty paragraphs
        if not para.text.strip():
            continue

        # Build paragraph HTML with formatting
        para_html_parts = []

        for run in para.runs:
            # Check for footnote reference
            if run.element.find(f"{w_ns}footnoteReference") is not None:
                fn_ref = run.element.find(f"{w_ns}footnoteReference")
                fn_id = fn_ref.get(qn('w:id'))

                if fn_id and fn_id not in ('0', '1'):
                    footnote_refs[fn_id] = True
                    # Create bidirectional link to footnote
                    para_html_parts.append(f'<sup><a href="#fn{fn_id}">{fn_id}</a></sup>')
                continue

            # Regular text run
            text = run.text
            if not text or not text.strip():
                continue

            # Strip URLs from text (remove hyperlinks)
            text = re.sub(r'\s*https?://\S+', '', text)

            # Escape HTML
            text = escape_html(text)

            # Strip hyperlinks from URL attributes (legacy)
            text = strip_hyperlinks_from_text(text)

            # Apply formatting - but be selective
            # Skip if all flags are set (indicates formatting artifact)
            rpr = run.element.find(f"{w_ns}rPr")
            has_all_flags = False
            if rpr is not None:
                has_b = rpr.find(f"{w_ns}b") is not None
                has_i = rpr.find(f"{w_ns}i") is not None
                has_u = rpr.find(f"{w_ns}u") is not None
                has_strike = rpr.find(f"{w_ns}strike") is not None

                # If all four are set, it's likely a formatting artifact - skip them
                if has_b and has_i and has_u and has_strike:
                    has_all_flags = True

            if rpr is not None and not has_all_flags:
                # Bold
                if rpr.find(f"{w_ns}b") is not None:
                    text = f"<strong>{text}</strong>"

                # Italic
                if rpr.find(f"{w_ns}i") is not None:
                    text = f"<em>{text}</em>"

                # Underline
                if rpr.find(f"{w_ns}u") is not None:
                    text = f"<u>{text}</u>"

                # Superscript
                va = rpr.find(f"{w_ns}vertAlign")
                if va is not None and va.get(qn('w:val')) == 'superscript':
                    text = f"<sup>{text}</sup>"

                # Subscript
                if va is not None and va.get(qn('w:val')) == 'subscript':
                    text = f"<sub>{text}</sub>"

                # Strike - only if not ALL flags are set
                if rpr.find(f"{w_ns}strike") is not None:
                    text = f"<s>{text}</s>"

            para_html_parts.append(text)

        # Check paragraph indentation for blockquote
        pf = para.paragraph_format
        is_blockquote = False
        if pf.left_indent and hasattr(pf.left_indent, 'inches'):
            if pf.left_indent.inches >= 0.2:
                is_blockquote = True

        para_html = ''.join(para_html_parts).strip()

        if para_html:
            if is_blockquote:
                html_parts.append(f"<blockquote>{para_html}</blockquote>")
            else:
                html_parts.append(f"<p>{para_html}</p>")

    # Build full opinion HTML and italicize case names
    opinion_html = '\n'.join(html_parts)
    opinion_html = italicize_case_names(opinion_html)

    # Filter footnotes to only those referenced in text, and italicize case names
    filtered_footnotes = {
        fn_id: italicize_case_names(footnotes_dict[fn_id])
        for fn_id in sorted(footnote_refs.keys())
        if fn_id in footnotes_dict
    }

    return opinion_html, filtered_footnotes


def strip_hyperlinks_from_text(text: str) -> str:
    """Remove hyperlink URLs and convert to plain text."""
    # This is basic - DOCX hyperlinks are in XML structure, not inline text
    # But check for common URL patterns just in case
    text = re.sub(r'https?://\S+', '', text)
    return text


def escape_html(text: str) -> str:
    """Escape HTML special characters."""
    replacements = {
        '&': '&amp;',
        '<': '&lt;',
        '>': '&gt;',
        '"': '&quot;',
        "'": '&#39;',
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text


def process_docx_file(docx_path: str) -> Tuple[str, Dict[str, str]]:
    """Process a single DOCX file and return formatted opinion + footnotes."""
    doc = Document(docx_path)

    # Extract footnotes first
    footnotes_dict = extract_footnotes(doc)

    # Extract opinion text with formatting
    opinion_html, filtered_footnotes = extract_formatted_opinion(doc, footnotes_dict)

    return opinion_html, filtered_footnotes


def verify_no_hyperlinks(opinion_html: str, footnotes: Dict[str, str]) -> bool:
    """Verify that opinion text contains no hyperlinks (except footnote backlinks)."""
    full_text = opinion_html + ' '.join(footnotes.values())

    # Check for common hyperlink patterns (excluding footnote links which are allowed)
    # Allow href="#fn" and href="#fnref" which are for bidirectional footnotes
    if 'href=' in full_text:
        # Check if ALL href= are for footnotes
        all_href = re.findall(r'href\s*=\s*["\']([^"\']*)["\']', full_text)
        for href in all_href:
            if not (href.startswith('#fn') or href.startswith('#fnref')):
                return False

    # Check for URLs outside of href attributes
    patterns = [
        r'https?://[^\s<>]',
    ]

    for pattern in patterns:
        if re.search(pattern, full_text):
            return False

    return True


if __name__ == '__main__':
    # Test with Woods case
    test_path = r'raw-opinions/Woods v. State_361 Ga. App. 844.Docx'

    opinion_html, footnotes = process_docx_file(test_path)

    print("=== OPINION HTML (first 1500 chars) ===")
    print(opinion_html[:1500])

    print(f"\n=== FOOTNOTES ({len(footnotes)} total) ===")
    for fn_id, fn_text in list(footnotes.items())[:3]:
        print(f"\nFootnote {fn_id}:")
        print(fn_text[:200])

    print(f"\n=== NO HYPERLINKS CHECK ===")
    has_no_links = verify_no_hyperlinks(opinion_html, footnotes)
    print(f"Clean (no hyperlinks): {has_no_links}")

    print(f"\n=== TOTAL LENGTH ===")
    print(f"Opinion HTML: {len(opinion_html)} chars")
    print(f"Footnotes: {sum(len(f) for f in footnotes.values())} chars")
